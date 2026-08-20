#!/usr/bin/env python3
"""
WSG markdown -> docx builder, Aug 2026 formatting spec.

Structure per wsg-article-formatting skill:
- Title style: article title only (first line).
- Heading 2 style: every section heading (H2). No H1/H3 styles in body.
- H3s in markdown become full-bold Normal paragraphs.
- One empty Normal paragraph between blocks; heading tight against the body below it.
- Lists: items tight; one empty line before first and after last.
- Tight line-groups (consecutive non-blank lines in one md paragraph) stay tight (say/don't triplets).

Visuals per style-guide/formatting-standards.md: Arial throughout; Title 24pt bold black;
Heading 2 18pt bold black; body 11pt black; links #0563C1 underlined; 2.0 line spacing; 0 space before/after.
"""
import re, sys
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Arial"
BLACK = RGBColor(0, 0, 0)
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)

def set_par_format(p):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

def style_run(run, size, bold=False, underline=False, color=BLACK):
    run.font.name = FONT
    r = run._element.rPr.rFonts if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
    run._element.get_or_add_rPr()
    rFonts = run._element.rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), FONT); rFonts.set(qn('w:hAnsi'), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    run.font.color.rgb = color

def add_hyperlink(paragraph, text, url, size, bold=False):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts'); rFonts.set(qn('w:ascii'), FONT); rFonts.set(qn('w:hAnsi'), FONT); rPr.append(rFonts)
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size * 2))); rPr.append(sz)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

TOKEN = re.compile(r'(\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|\*([^*]+)\*)')

def add_inline(p, text, size, base_bold=False):
    pos = 0
    for m in TOKEN.finditer(text):
        if m.start() > pos:
            style_run(p.add_run(text[pos:m.start()]), size, bold=base_bold)
        if m.group(2) is not None:
            add_hyperlink(p, m.group(2), m.group(3), size, bold=base_bold)
        elif m.group(4) is not None:
            style_run(p.add_run(m.group(4)), size, bold=True)
        else:
            r = p.add_run(m.group(5)); style_run(r, size, bold=base_bold); r.font.italic = True
        pos = m.end()
    if pos < len(text):
        style_run(p.add_run(text[pos:]), size, bold=base_bold)

def main(md_path, out_path):
    raw = open(md_path, encoding="utf-8").read()
    m = re.match(r"^---\n.*?\n---\n", raw, flags=re.DOTALL)
    body = raw[m.end():] if m else raw

    doc = Document()
    # Normal style base
    normal = doc.styles['Normal']
    normal.font.name = FONT; normal.font.size = Pt(11); normal.font.color.rgb = BLACK

    def empty_par():
        p = doc.add_paragraph(); set_par_format(p); return p

    # Split into blocks on blank lines
    blocks = [b for b in re.split(r'\n\s*\n', body.strip()) if b.strip()]
    first_content = True  # nothing above Title
    prev_was_heading = False

    for block in blocks:
        lines = [l.rstrip() for l in block.split('\n') if l.strip()]
        if not lines:
            continue
        first = lines[0]

        if first.startswith('# ') and not first.startswith('## '):
            # Title
            p = doc.add_paragraph(style=doc.styles['Title'])
            set_par_format(p)
            add_inline(p, first[2:].strip(), 24)
            for r in p.runs: r.font.bold = True; r.font.color.rgb = BLACK; r.font.name = FONT; r.font.size = Pt(24)
            # kill Title style bottom border artifact
            pPr = p._p.get_or_add_pPr()
            for bdr in pPr.findall(qn('w:pBdr')): pPr.remove(bdr)
            first_content = False
            prev_was_heading = True
            continue

        if first.startswith('## '):
            if not first_content:
                empty_par()
            p = doc.add_paragraph(style=doc.styles['Heading 2'])
            set_par_format(p)
            add_inline(p, first[3:].strip(), 18)
            for r in p.runs: r.font.bold = True; r.font.color.rgb = BLACK; r.font.name = FONT; r.font.size = Pt(18)
            first_content = False
            prev_was_heading = True
            continue

        if not first_content and not prev_was_heading:
            empty_par()
        first_content = False
        prev_was_heading = False

        if first.startswith('### '):
            # bold body paragraph(s)
            p = doc.add_paragraph(); set_par_format(p)
            add_inline(p, re.sub(r'\*\*', '', first[4:].strip()), 11, base_bold=True)
            # any continuation lines in same block -> tight lines after
            for ln in lines[1:]:
                p2 = doc.add_paragraph(); set_par_format(p2)
                add_inline(p2, ln, 11)
            continue

        if all(re.match(r'^[-*] ', l) for l in lines):
            for ln in lines:
                p = doc.add_paragraph(style=doc.styles['List Bullet'])
                set_par_format(p)
                add_inline(p, ln[2:].strip(), 11)
                for r in p.runs:
                    if r.font.size is None: style_run(r, 11)
            continue

        if all(re.match(r'^\d+\. ', l) for l in lines):
            for ln in lines:
                p = doc.add_paragraph(style=doc.styles['List Number'])
                set_par_format(p)
                add_inline(p, re.sub(r'^\d+\. ', '', ln), 11)
            continue

        # plain block: consecutive lines stay tight (say/don't triplets etc.)
        for ln in lines:
            ln = re.sub(r'^> ?', '', ln)
            p = doc.add_paragraph(); set_par_format(p)
            add_inline(p, ln, 11)

    doc.save(out_path)
    print(f"built {out_path}")

if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
