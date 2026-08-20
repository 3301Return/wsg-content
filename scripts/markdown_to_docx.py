#!/usr/bin/env python3
"""
markdown_to_docx.py — build a Word doc from a WSG draft markdown file.

WSG FORMATTING STANDARD (locked May 24, 2026):
- Font: Arial throughout (body, headings, bullets).
- H1 title: 24pt, bold, black.
- H2 section: 18pt, bold, black.
- H3 firm/sub: 14pt, bold, underlined, black.
- Body: 11pt, black.
- Line spacing: 2.0 (double) on EVERYTHING.
- Inline links render as real Word hyperlinks (blue, underlined).
- Pre-build integrity guard added June 2026 to catch silent Edit-tool truncation.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


FONT_NAME = "Arial"
BODY_PT = 11
H1_PT = 24
H2_PT = 18
H3_PT = 14
BLACK = RGBColor(0x00, 0x00, 0x00)
LINK_BLUE = "0563C1"

LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def strip_frontmatter(text):
    if not text.startswith("---"):
        return text
    parts = text.split("---", 2)
    if len(parts) < 3:
        return text
    return parts[2].lstrip("\n")


def _set_run_font(run, size_pt, bold=False, underline=False):
    run.font.name = FONT_NAME
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = BLACK
    run.bold = bold
    run.underline = underline


def _set_paragraph_double_spacing(paragraph):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def add_hyperlink(paragraph, url, text, size_pt, bold=False, underline=True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size_pt * 2))
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    rPr.append(color)

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def render_inline(paragraph, text, size_pt, base_bold=False, base_underline=False):
    pos = 0
    while pos < len(text):
        link_m = LINK_RE.search(text, pos)
        bold_m = BOLD_RE.search(text, pos)

        next_m = None
        kind = None
        if link_m and (not bold_m or link_m.start() < bold_m.start()):
            next_m, kind = link_m, "link"
        elif bold_m:
            next_m, kind = bold_m, "bold"

        if not next_m:
            chunk = text[pos:]
            if chunk:
                run = paragraph.add_run(chunk)
                _set_run_font(run, size_pt, bold=base_bold, underline=base_underline)
            return

        if next_m.start() > pos:
            chunk = text[pos:next_m.start()]
            run = paragraph.add_run(chunk)
            _set_run_font(run, size_pt, bold=base_bold, underline=base_underline)

        if kind == "link":
            label, url = next_m.group(1), next_m.group(2)
            add_hyperlink(paragraph, url, label, size_pt, bold=base_bold, underline=True)
        else:
            run = paragraph.add_run(next_m.group(1))
            _set_run_font(run, size_pt, bold=True, underline=base_underline)

        pos = next_m.end()


def _apply_normal_defaults(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = BLACK
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT_NAME)


def _verify_markdown_integrity(body, md_path):
    """Pre-build integrity checks. Catches the Edit-tool truncation footprint."""
    errors = []

    if body.count("**") % 2 != 0:
        errors.append("Unbalanced bold markers (odd count of **).")

    if body.count("[") != body.count("]"):
        errors.append(f"Unbalanced brackets: [={body.count('[')} ]={body.count(']')}.")

    nonempty = [ln.rstrip() for ln in body.splitlines() if ln.strip()]
    if nonempty:
        last = nonempty[-1]
        if re.match(r"^(\d+\.\s|[-*]\s)", last):
            stripped = re.sub(r"^(\d+\.\s|[-*]\s)", "", last).strip()
            if re.fullmatch(r"\*\*[^*]+\*\*\s*", stripped):
                errors.append(
                    "File ends with a list item containing only a bolded sentence. "
                    "This is the Edit-tool truncation signature."
                )

    raw_lines = body.splitlines()
    while raw_lines and not raw_lines[-1].strip():
        raw_lines.pop()
    if raw_lines and raw_lines[-1] != raw_lines[-1].rstrip():
        errors.append("Last content line has trailing whitespace (partial-Edit footprint).")

    if errors:
        print("BUILD ABORTED for " + str(md_path), file=sys.stderr)
        for e in errors:
            print("  X  " + e, file=sys.stderr)
        sys.exit(3)


def build(md_path, out_path):
    raw = Path(md_path).read_text(encoding="utf-8")
    body = strip_frontmatter(raw)

    _verify_markdown_integrity(body, md_path)

    doc = Document()
    _apply_normal_defaults(doc)

    for line in body.splitlines():
        line = line.rstrip()
        if not line.strip():
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            render_inline(p, line[4:].strip(), H3_PT, base_bold=True, base_underline=True)
            _set_paragraph_double_spacing(p)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            render_inline(p, line[3:].strip(), H2_PT, base_bold=True)
            _set_paragraph_double_spacing(p)
        elif line.startswith("# "):
            p = doc.add_paragraph()
            render_inline(p, line[2:].strip(), H1_PT, base_bold=True)
            _set_paragraph_double_spacing(p)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            render_inline(p, line[2:].strip(), BODY_PT)
            _set_paragraph_double_spacing(p)
        elif re.match(r"^\d+\.\s", line):
            content = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            render_inline(p, content.strip(), BODY_PT)
            _set_paragraph_double_spacing(p)
        else:
            p = doc.add_paragraph()
            render_inline(p, line, BODY_PT)
            _set_paragraph_double_spacing(p)

    # POST-BUILD VERIFICATION: ensure the docx's last paragraph matches the markdown's last
    # non-empty line. If the parser silently truncated mid-document, abort instead of saving.
    md_nonempty = [ln.rstrip() for ln in body.splitlines() if ln.strip()]
    md_last_clean = re.sub(r"\*\*|\[([^\]]+)\]\([^)]+\)", lambda m: m.group(1) if m.group(0).startswith("[") else "", md_nonempty[-1])
    md_last_clean = md_last_clean.strip()
    docx_last = ""
    for p in doc.paragraphs:
        if p.text.strip():
            docx_last = p.text.strip()
    # compare last 40 chars after stripping markdown markers
    md_tail = md_last_clean[-40:].lower().replace(" ", "")
    docx_tail = docx_last[-60:].lower().replace(" ", "")
    if md_tail and md_tail not in docx_tail:
        print(f"BUILD ABORTED: docx tail does not match markdown tail.", file=sys.stderr)
        print(f"  md last:  {md_last_clean[-80:]!r}", file=sys.stderr)
        print(f"  docx last: {docx_last[-80:]!r}", file=sys.stderr)
        sys.exit(4)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print("OK  " + str(md_path) + " -> " + str(out_path))


def main():
    if len(sys.argv) != 3:
        print("Usage: markdown_to_docx.py <input.md> <output.docx>", file=sys.stderr)
        sys.exit(2)
    md_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    if not md_path.exists():
        print("ERR  no such file: " + str(md_path), file=sys.stderr)
        sys.exit(1)
    build(md_path, out_path)


if __name__ == "__main__":
    main()
